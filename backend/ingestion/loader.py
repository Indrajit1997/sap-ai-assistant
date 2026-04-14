from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from urllib.parse import urlparse

import structlog

log = structlog.get_logger()


@dataclass
class Document:
    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.sha256(self.content.encode()).hexdigest()[:16]


class DocumentLoader:
    """Loads documents from various formats into a unified Document structure."""

    SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".md", ".txt", ".docx"}

    def load_file(self, path: Path) -> list[Document]:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}")

        log.info("loading_document", path=str(path), type=ext)

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext in (".html", ".htm"):
            return self._load_html(path)
        elif ext == ".md":
            return self._load_markdown(path)
        elif ext == ".txt":
            return self._load_text(path)
        elif ext == ".docx":
            return self._load_docx(path)
        return []

    def load_directory(self, dir_path: Path, recursive: bool = True) -> list[Document]:
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        documents = []
        pattern = "**/*" if recursive else "*"
        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    documents.extend(self.load_file(file_path))
                except Exception as e:
                    log.warning("skip_file", path=str(file_path), error=str(e))
        log.info("directory_loaded", path=str(dir_path), doc_count=len(documents))
        return documents

    def load_url(self, url: str) -> list[Document]:
        """Fetch a public URL and extract its text content."""
        import httpx
        from bs4 import BeautifulSoup

        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            raise ValueError(f"Invalid URL scheme: {parsed.scheme}. Only http/https supported.")

        log.info("loading_url", url=url)

        response = httpx.get(
            url,
            follow_redirects=True,
            timeout=30,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.5",
            },
        )
        response.raise_for_status()

        content_type = response.headers.get("content-type", "")

        if "text/html" in content_type or "application/xhtml" in content_type:
            soup = BeautifulSoup(response.text, "lxml")
            title = soup.title.string.strip() if soup.title and soup.title.string else ""

            for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
                tag.decompose()

            # Try to find main content area
            main = soup.find("main") or soup.find("article") or soup.find("div", {"role": "main"}) or soup.body
            text = main.get_text(separator="\n", strip=True) if main else ""
        else:
            # Plain text or other
            text = response.text
            title = ""

        if not text or not text.strip():
            log.warning("url_empty_content", url=url)
            return []

        return [Document(
            content=text.strip(),
            metadata={"source": url, "type": "url", "title": title},
        )]

    def load_urls(self, urls: list[str]) -> list[Document]:
        """Fetch multiple URLs and return all documents."""
        documents = []
        for url in urls:
            try:
                documents.extend(self.load_url(url.strip()))
            except Exception as e:
                log.warning("skip_url", url=url, error=str(e))
        log.info("urls_loaded", total_urls=len(urls), doc_count=len(documents))
        return documents

    def _load_pdf(self, path: Path) -> list[Document]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        documents = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                documents.append(Document(
                    content=text.strip(),
                    metadata={"source": str(path), "page": i + 1, "type": "pdf"},
                ))
        return documents

    def _load_html(self, path: Path) -> list[Document]:
        from bs4 import BeautifulSoup

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "lxml")

        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        if text:
            return [Document(
                content=text,
                metadata={"source": str(path), "type": "html", "title": soup.title.string if soup.title else ""},
            )]
        return []

    def _load_markdown(self, path: Path) -> list[Document]:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        if content.strip():
            return [Document(
                content=content.strip(),
                metadata={"source": str(path), "type": "markdown"},
            )]
        return []

    def _load_text(self, path: Path) -> list[Document]:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        if content.strip():
            return [Document(
                content=content.strip(),
                metadata={"source": str(path), "type": "text"},
            )]
        return []

    def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        if text:
            return [Document(
                content=text,
                metadata={"source": str(path), "type": "docx"},
            )]
        return []
